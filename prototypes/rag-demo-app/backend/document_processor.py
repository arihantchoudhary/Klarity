from typing import List, Any, Tuple, Dict
import os
import pdfplumber
from pdf2image import convert_from_path
import io
import base64
from PIL import Image
from tqdm import tqdm
import streamlit as st
import logging
from pathlib import Path
import fitz
import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles the processing of PDF documents, extracting text, tables, and images."""
    
    def __init__(self, output_path: str):
        """Initialize the document processor.
        
        Args:
            output_path: Directory to save extracted content
        """
        self.output_path = Path(output_path)
        self.output_path.mkdir(parents=True, exist_ok=True)
    
    def process_pdf(self, file_path: str) -> Tuple[List[str], List[pd.DataFrame], List[Image.Image]]:
        """Process a PDF file and extract text, tables, and images.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (texts, tables, images)
        """
        texts = []
        tables = []
        images = []
        
        try:
            with fitz.open(file_path) as pdf:
                for page_num, page in enumerate(pdf):
                    # Extract text
                    text = page.get_text()
                    if text.strip():
                        texts.append(text)
                    
                    # Extract images
                    for img_index, img in enumerate(page.get_images()):
                        try:
                            xref = img[0]
                            base_image = pdf.extract_image(xref)
                            image_bytes = base_image["image"]
                            image = Image.open(io.BytesIO(image_bytes))
                            images.append(image)
                        except Exception as e:
                            logger.warning(f"Failed to extract image {img_index} from page {page_num}: {e}")
                    
                    # Extract tables (basic implementation)
                    # Note: This is a simplified approach. For better table extraction,
                    # consider using specialized libraries like tabula-py
                    tables_text = page.get_text("blocks")
                    if tables_text:
                        try:
                            df = pd.DataFrame([t[4] for t in tables_text if len(t) > 4])
                            if not df.empty:
                                tables.append(df)
                        except Exception as e:
                            logger.warning(f"Failed to extract table from page {page_num}: {e}")
        
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
        
        return texts, tables, images
    
    def process_docx(self, file_path: str) -> Tuple[List[str], List[pd.DataFrame], List[Image.Image]]:
        """Process a Word document and extract text, tables, and images.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Tuple of (texts, tables, images)
        """
        texts = []
        tables = []
        images = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    data.append(row_data)
                if data:
                    tables.append(pd.DataFrame(data[1:], columns=data[0]))
            
            # Note: Image extraction from DOCX is more complex and might require
            # additional processing of doc.inline_shapes or doc.part.related_parts
        
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise
        
        return texts, tables, images
    
    def process_text(self, file_path: str) -> List[str]:
        """Process a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            List of text chunks
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return [text] if text.strip() else []
        
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    def save_extracted_content(self, file_id: str, texts: List[str], tables: List[pd.DataFrame], 
                             images: List[Image.Image]) -> Dict[str, List[str]]:
        """Save extracted content to files.
        
        Args:
            file_id: Unique identifier for the document
            texts: List of extracted text chunks
            tables: List of extracted tables
            images: List of extracted images
            
        Returns:
            Dictionary with paths to saved content
        """
        saved_paths = {
            'texts': [],
            'tables': [],
            'images': []
        }
        
        # Save texts
        for i, text in enumerate(texts):
            text_path = self.output_path / f"{file_id}_text_{i}.txt"
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text)
            saved_paths['texts'].append(str(text_path))
        
        # Save tables
        for i, table in enumerate(tables):
            table_path = self.output_path / f"{file_id}_table_{i}.csv"
            table.to_csv(table_path, index=False)
            saved_paths['tables'].append(str(table_path))
        
        # Save images
        for i, image in enumerate(images):
            image_path = self.output_path / f"{file_id}_image_{i}.png"
            image.save(image_path, 'PNG')
            saved_paths['images'].append(str(image_path))
        
        return saved_paths

def main():
    """Streamlit frontend for PDF processing"""
    st.title("PDF Document Processor")
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # File uploader
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
    if uploaded_file is not None:
        # Save uploaded file temporarily
        with open("temp.pdf", "wb") as f:
            f.write(uploaded_file.getvalue())
        
        # Process the file
        with st.spinner('Processing PDF...'):
            texts, tables, images = processor.process_pdf("temp.pdf")
        
        # Display results
        st.success("Processing complete!")
        
        # Show text content
        if texts:
            st.subheader("Extracted Text")
            for i, text in enumerate(texts):
                with st.expander(f"Text chunk {i+1}"):
                    st.write(text)
        
        # Show tables
        if tables:
            st.subheader("Extracted Tables")
            for i, table in enumerate(tables):
                with st.expander(f"Table {i+1}"):
                    st.write(table)
        
        # Show images
        if images:
            st.subheader("Extracted Images")
            for i, image in enumerate(images):
                st.image(image, caption=f"Image {i+1}")
        
        # Cleanup
        os.remove("temp.pdf")
    
    # Directory processing
    st.subheader("Process Directory")
    dir_path = st.text_input("Enter directory path containing PDFs")
    if st.button("Process Directory") and dir_path:
        if os.path.exists(dir_path):
            with st.spinner('Processing directory...'):
                results = processor.process_directory(dir_path)
            
            st.success("Directory processing complete!")
            for filename, (texts, tables, images) in results.items():
                with st.expander(f"File: {filename}"):
                    st.write(f"Number of text chunks: {len(texts)}")
                    st.write(f"Number of tables: {len(tables)}")
                    st.write(f"Number of images: {len(images)}")
        else:
            st.error("Directory not found!")

if __name__ == "__main__":
    main()